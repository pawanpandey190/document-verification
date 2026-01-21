import pandas as pd
from docx import Document
import os

TEMPLATE_PATH = "/Users/pawanpandey/Downloads/CondtionalAcceptanceLetter_Foundation_IFP_19.01.26.docx"
EXCEL_PATH = "/Users/pawanpandey/Documents/document-validation/student_validation_report.xlsx"
OUTPUT_DIR = "generated_letters"

os.makedirs(OUTPUT_DIR, exist_ok=True)


from docx.shared import RGBColor

def safe_replace_placeholders(doc, replacements: dict):
    def process_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)

        replaced = full_text
        for key, value in replacements.items():
            replaced = replaced.replace(key, value)

        if replaced != full_text:
            # Save paragraph style
            p_style = paragraph.style

            # Clear runs safely
            for run in paragraph.runs:
                run.text = ""

            # Create ONE new run
            run = paragraph.add_run(replaced)
            run.bold = False
            run.italic = False
            run.font.color.rgb = RGBColor(0, 0, 0)

            # Restore paragraph style
            paragraph.style = p_style

    # Body
    for paragraph in doc.paragraphs:
        process_paragraph(paragraph)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_paragraph(paragraph)

    # Headers & Footers
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            process_paragraph(paragraph)
        for paragraph in section.footer.paragraphs:
            process_paragraph(paragraph)

def generate_letters(EXCEL_PATH):
    df = pd.read_excel(EXCEL_PATH)

    for _, row in df.iterrows():
        doc = Document(TEMPLATE_PATH)

        replacements = {
            "{{Student_name}}": str(row["Passport Holder Name"])
            
            
        }

        safe_replace_placeholders(doc, replacements)

        output_file = f"Conditional_Offer_{row['Passport Holder Name'].replace(' ', '_')}.docx"
        output_path = os.path.join(OUTPUT_DIR, output_file)

        doc.save(output_path)

    print("✅ All letters generated successfully!")



